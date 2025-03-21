from flask import Flask, render_template, request, send_file, send_from_directory
from flask_cors import CORS
import os
os.environ["TESSERACT_PATH"] = "/usr/bin/tesseract"
os.environ["PATH"] += os.pathsep + os.environ["TESSERACT_PATH"]
print(f"System PATH: {os.environ.get('PATH')}")
print(f"Tesseract PATH: {os.getenv('TESSERACT_PATH')}")

from pdf2docx import Converter
from PIL import Image
import pytesseract
import pandas as pd
from docx import Document
from pdf2image import convert_from_path
from datetime import datetime, timedelta
from threading import Thread
import shutil
import time
import subprocess
import shutil
from concurrent.futures import ThreadPoolExecutor
import multiprocessing

def check_tesseract():
    try:
        # Try multiple ways to find tesseract
        tesseract_path = shutil.which('tesseract')
        if tesseract_path:
            print(f"Found tesseract at: {tesseract_path}")
        else:
            print("Tesseract not found in PATH")
            
        # Try running tesseract
        result = subprocess.run(['tesseract', '--version'], 
                              capture_output=True, 
                              text=True)
        print(f"Tesseract version output: {result.stdout}")
        return True
    except Exception as e:
        print(f"Detailed Tesseract error: {str(e)}")
        print(f"Current PATH: {os.environ.get('PATH', 'Not set')}")
        return False

# Add this near the start of your app
print("Checking Tesseract installation...")
check_tesseract()

pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_PATH', '/usr/bin/tesseract')


app = Flask(__name__)
CORS(app, resources={
    r"/upload": {"origins": ["https://www.simplepdftoword.com"]},
    r"/system-check": {"origins": ["https://www.simplepdftoword.com"]},
    r"/sitemap.xml": {"origins": ["https://www.simplepdftoword.com"]},
    r"/robots.txt": {"origins": ["https://www.simplepdftoword.com"]}
})

# Create folders
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER

ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files():
    """Remove files older than 1 minute"""
    try:
        current_time = datetime.now()
        for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    filepath = os.path.join(folder, filename)
                    file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if current_time - file_time > timedelta(minutes=2):  # Changed from 10 to 2 minute
                        try:
                            if os.path.isfile(filepath):
                                os.remove(filepath)
                            elif os.path.isdir(filepath):
                                shutil.rmtree(filepath)
                        except Exception as e:
                            print(f"Error removing {filepath}: {str(e)}")
    except Exception as e:
        print(f"Error in cleanup: {str(e)}")

def periodic_cleanup():
    while True:
        cleanup_old_files()
        time.sleep(120)  # Changed from 300 to 120 seconds to check more frequently

def process_page(page_image):
    return pytesseract.image_to_string(page_image)

def convert_pdf_to_word(pdf_path, output_path):
    try:
        # Convert PDF to images
        images = convert_from_path(pdf_path)
        doc = Document()
        
        # Use number of CPU cores for parallel processing
        num_workers = multiprocessing.cpu_count()
        
        # Process pages in parallel
        with ThreadPoolExecutor(max_workers=num_workers) as executor:
            text_results = list(executor.map(process_page, images))
        
        # Add processed text to document
        for i, text in enumerate(text_results):
            if text.strip():
                doc.add_paragraph(text)
            # Only add page break if not the last page
            if i < len(text_results) - 1:
                doc.add_page_break()
        
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"PDF to Word conversion failed: {str(e)}")

def convert_pdf_to_excel(pdf_path, output_path):
    try:
        # Convert PDF to images
        images = convert_from_path(pdf_path)
        
        # Extract text from all pages
        all_text = []
        for image in images:
            text = pytesseract.image_to_string(image)
            all_text.append(text)
        
        # Combine text and convert to Excel
        combined_text = '\n'.join(all_text)
        df = pd.DataFrame([line.split() for line in combined_text.split('\n') if line.strip()])
        df.to_excel(output_path, index=False)
    except Exception as e:
        raise Exception(f"PDF to Excel conversion failed: {str(e)}")

def convert_image_to_excel(image_path, output_path):
    # Extract text from image using OCR
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)
    # Convert to Excel
    df = pd.DataFrame([line.split() for line in text.split('\n')])
    df.to_excel(output_path, index=False)

def convert_image_to_word(image_path, output_path):
    # Extract text from image using OCR
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)
    # Create Word document
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    cleanup_old_files()  # Run cleanup before processing new file
    # Define max file size (20MB)
    MAX_FILE_SIZE = 200 * 1024 * 1024  # 20MB in bytes
    
    if 'file' not in request.files:
        return 'Please select a file', 400
    
    file = request.files['file']
    output_format = request.form.get('output_format', 'docx')
    
    if file.filename == '':
        return 'No file selected', 400

    # Check file size
    file.seek(0, 2)  # Seek to end of file
    file_size = file.tell()  # Get current position (file size)
    file.seek(0)  # Reset file pointer to beginning
    
    if file_size > MAX_FILE_SIZE:
        return 'File size exceeds 200MB limit', 400

    if file and allowed_file(file.filename):
        # Save original file
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(input_path)
        
        # Generate output filename
        filename_without_ext = os.path.splitext(file.filename)[0]
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{filename_without_ext}.{output_format}")
        
        # Convert based on input type and desired output
        try:
            file_ext = file.filename.rsplit('.', 1)[1].lower()
            
            if file_ext == 'pdf':
                if output_format == 'docx':
                    convert_pdf_to_word(input_path, output_path)
                elif output_format in ['xlsx', 'csv']:
                    convert_pdf_to_excel(input_path, output_path)
            elif file_ext in ['png', 'jpg', 'jpeg', 'gif']:
                if output_format == 'docx':
                    convert_image_to_word(input_path, output_path)
                elif output_format in ['xlsx', 'csv']:
                    convert_image_to_excel(input_path, output_path)
            
            return send_file(output_path, as_attachment=True)
            
        except Exception as e:
            return f'Conversion error: {str(e)}', 500
            
    return 'Unsupported file type', 400

@app.route('/upload', methods=['GET'])
def upload_get():
    # Return a proper response for GET requests to /upload
    #return render_template('index.html')  # Redirect to home page
    # Or alternatively:
     return 'This endpoint only accepts file uploads via POST requests', 200

@app.route('/system-check')
def system_check():
    try:
        # Check tesseract
        tesseract_version = subprocess.run(['tesseract', '--version'], 
                                         capture_output=True, 
                                         text=True)
        
        # Check poppler
        poppler_version = subprocess.run(['pdftoppm', '-v'], 
                                       capture_output=True, 
                                       text=True, 
                                       stderr=subprocess.STDOUT)
        
        return {
            'tesseract_version': tesseract_version.stdout,
            'poppler_version': poppler_version.stdout,
            'PATH': os.environ.get('PATH', 'Not set'),
            'TESSERACT_PATH': os.environ.get('TESSERACT_PATH', 'Not set')
        }
    except Exception as e:
        return {'error': str(e)}

@app.route('/sitemap.xml')
def serve_sitemap():
    return send_from_directory('templates', 'sitemap.xml')

@app.route('/robots.txt')
def serve_robots():
    return send_from_directory('templates', 'robots.txt')

if __name__ == '__main__':
    # Start cleanup thread
    cleanup_thread = Thread(target=periodic_cleanup, daemon=True)
    cleanup_thread.start()
    
    # Your existing app.run line
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)